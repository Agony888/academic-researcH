"""Markdown and native Word report generation."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import AppConfig
from utils import DailyReport, Paper, ReportItem, format_date_cn, write_text_file


def generate_markdown(report: DailyReport, config: AppConfig) -> Path:
    path = config.daily_dir / f"{report.report_date.isoformat()}.md"
    content = build_markdown(report)
    write_text_file(path, content)
    report.markdown_path = path
    return path


def generate_word(report: DailyReport, config: AppConfig) -> Path:
    """Generate a native .docx Word document with styled text.

    Earlier versions emitted HTML with a Word-like extension. This function now
    creates a real Office Open XML document so Word/WPS opens a formatted report
    directly as text, tables and headings.
    """

    path = config.daily_dir / f"{report.report_date.isoformat()}.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _setup_document(document)
    _add_cover(document, report)
    _add_hotspot_summary(document, report)
    _add_toc(document, report)
    _add_papers(document, report)
    _add_stats(document, report)
    document.save(path)
    report.word_path = path
    return path


def generate_docx_from_markdown(markdown_path: str | Path, config: AppConfig) -> Path:
    """Rebuild a native DOCX from a UTF-8 Markdown report.

    This is useful when a report markdown is correct but a document was created
    through a console path with a legacy code page. The function reads the file
    bytes as UTF-8 directly, so Chinese text does not pass through the terminal.
    """

    source = Path(markdown_path)
    if not source.is_absolute():
        source = config.project_root / source
    content = source.read_text(encoding="utf-8")
    output = source.with_suffix(".docx")

    document = Document()
    _setup_document(document)

    in_code_block = False
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("\x60\x60\x60"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            document.add_paragraph(line)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(_strip_markdown(line[2:]))
            run.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(31, 84, 150)
        elif line.startswith("## "):
            document.add_heading(_strip_markdown(line[3:]), level=1)
        elif line.startswith("### "):
            title = re.sub(r"<a\s+id=.*?</a>", "", line[4:])
            document.add_heading(_strip_markdown(title), level=2)
        elif line.startswith("#### "):
            document.add_heading(_strip_markdown(line[5:]), level=3)
        elif line.startswith(">"):
            _add_quote_block(document, _strip_markdown(line.lstrip("> ").strip()))
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(_strip_markdown(line), style="List Number")
        elif line.startswith("- "):
            document.add_paragraph(_strip_markdown(line[2:]), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown(line))

    document.save(output)
    return output


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, RGBColor(31, 84, 150)),
        ("Heading 1", 17, RGBColor(31, 84, 150)),
        ("Heading 2", 14, RGBColor(47, 93, 154)),
        ("Heading 3", 12, RGBColor(39, 73, 125)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def _add_cover(document: Document, report: DailyReport) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"每日SSCI文献简报｜{report.report_date.isoformat()}")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 84, 150)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("学前与基础教育 · AI教学应用 · 教师数字/智能/数据素养 · 教师教育")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(89, 104, 130)

    if report.notice:
        _add_quote_block(document, report.notice)

    table = document.add_table(rows=4, cols=2)
    table.style = "Light Shading Accent 1"
    rows = [
        ("生成时间", report.generated_at.strftime("%Y-%m-%d %H:%M:%S（北京时间）")),
        ("检索窗口", f"{report.window_start.strftime('%Y-%m-%d')} — {report.window_end.strftime('%Y-%m-%d')}"),
        ("检索与筛选", f"检索 {report.total_found} 篇；筛选 {report.total_filtered} 篇；AI成功 {report.total_success} 篇"),
        ("本地文件", "Markdown 与 Word 文档已保存至 daily/ 目录"),
    ]
    for row, (key, value) in zip(table.rows, rows, strict=True):
        row.cells[0].text = key
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()


def _add_hotspot_summary(document: Document, report: DailyReport) -> None:
    document.add_heading("一、今日研究亮点总结", level=1)
    _add_quote_block(document, report.hotspot_summary.strip() or "今日未生成文献亮点总结。")


def _add_toc(document: Document, report: DailyReport) -> None:
    document.add_heading("目录", level=1)
    if not report.items:
        document.add_paragraph("今日未筛选到符合条件的论文。")
        return
    for index, item in enumerate(report.items, start=1):
        p = document.add_paragraph(style="List Number")
        p.add_run(f"{item.summary.chinese_title}").bold = True
        p.add_run(f"｜{item.paper.journal or '未知期刊'}｜{format_date_cn(item.paper.published_date)}")


def _add_papers(document: Document, report: DailyReport) -> None:
    if report.items:
        document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("二、文献摘要", level=1)
    for index, item in enumerate(report.items, start=1):
        if index > 1:
            document.add_page_break()
        _add_one_paper(document, index, item)


def _add_one_paper(document: Document, index: int, item: ReportItem) -> None:
    paper = item.paper
    summary = item.summary
    document.add_heading(f"{index}. {paper.title} / {summary.chinese_title}", level=2)

    table = document.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    _add_table_row(table, "标题", f"{paper.title} / {summary.chinese_title}")
    _add_table_row(table, "作者", summary.authors_text or (", ".join(paper.authors[:8]) if paper.authors else "摘要未说明"))
    _add_table_row(table, "期刊及影响因子", summary.journal_impact or _journal_metadata_text(paper))
    _add_table_row(table, "JIF/影响因子", paper.impact_factor or "白名单未提供")
    _add_table_row(table, "JCR分区", paper.quartile or "白名单未提供")
    _add_table_row(table, "CiteScore", paper.citescore or "白名单未提供")
    _add_table_row(table, "出版社", paper.publisher or "白名单未提供")
    _add_table_row(table, "卷期/DOI", summary.volume_issue_doi or paper.doi or "摘要未说明")
    _add_table_row(table, "在线发表时间", summary.online_date or format_date_cn(paper.published_date))
    _add_table_row(table, "SSCI分类", "；".join(paper.ssci_categories) if paper.ssci_categories else "SSCI白名单匹配；分类未提供")

    link = paper.url or (f"https://doi.org/{paper.doi}" if paper.doi else "")
    if link:
        row = table.add_row()
        row.cells[0].text = "论文链接"
        row.cells[0].paragraphs[0].runs[0].bold = True
        _add_hyperlink(row.cells[1].paragraphs[0], link, link)

    sections = [
        ("研究摘要", summary.research_abstract),
        ("研究主要内容", summary.main_content),
        ("研究论点", summary.argument),
        ("研究问题", summary.research_question),
        ("研究方法", summary.methods),
        ("主要发现", summary.findings),
        ("对教师教育研究的启示", summary.implications_teacher_education),
        ("APA引用格式", summary.apa_citation),
    ]
    for heading, text in sections:
        document.add_heading(heading, level=3)
        _add_text_block(document, text)


def _add_stats(document: Document, report: DailyReport) -> None:
    document.add_heading("运行统计", level=1)
    table = document.add_table(rows=5, cols=2)
    table.style = "Light Shading Accent 1"
    stats = [
        ("检索数量", str(report.total_found)),
        ("筛选数量", str(report.total_filtered)),
        ("AI成功数量", str(report.total_success)),
        ("失败数量", str(report.total_failed)),
        ("API耗时", f"{report.api_elapsed_seconds:.2f} 秒"),
    ]
    for row, (key, value) in zip(table.rows, stats, strict=True):
        row.cells[0].text = key
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True


def _add_table_row(table, key: str, value: str) -> None:
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True


def _add_quote_block(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(66, 82, 110)
    _shade_paragraph(p, "EEF4FF")


def _add_text_block(document: Document, text: str) -> None:
    cleaned = text.strip() or "摘要未说明。"
    lines = [line.strip() for line in cleaned.replace("；", "；\n").splitlines() if line.strip()]
    if not lines:
        document.add_paragraph("摘要未说明。")
        return
    for line in lines:
        if line.startswith(("-", "•", "1.", "2.", "3.", "①", "②", "③")):
            document.add_paragraph(line.lstrip("-• ").strip(), style="List Bullet")
        else:
            document.add_paragraph(line)


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _strip_markdown(text: str) -> str:
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1（\2）", cleaned)
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    return cleaned.strip()


def build_markdown(report: DailyReport) -> str:
    lines: list[str] = []
    lines.append(f"# 每日SSCI文献简报（{report.report_date.isoformat()}）")
    lines.append("")
    lines.append(f"> 生成时间：{report.generated_at.strftime('%Y-%m-%d %H:%M:%S')}  ")
    lines.append(f"> 检索窗口：{report.window_start.strftime('%Y-%m-%d %H:%M')} — {report.window_end.strftime('%Y-%m-%d %H:%M')}（北京时间）")
    if report.notice:
        lines.append(f"> ⚠️ {report.notice}")
    lines.append("")
    lines.append("## 一、今日研究亮点总结")
    lines.append("")
    lines.append(report.hotspot_summary.strip() or "今日未生成热点总结。")
    lines.append("")
    lines.append("## 目录")
    lines.append("")
    if report.items:
        for index, item in enumerate(report.items, start=1):
            lines.append(f"{index}. [{item.paper.title} / {item.summary.chinese_title}](#paper-{index})")
    else:
        lines.append("- 今日未筛选到符合条件的论文。")
    lines.append("")
    lines.append("## 二、文献摘要")
    lines.append("")
    for index, item in enumerate(report.items, start=1):
        lines.extend(_item_markdown(index, item))
    lines.append("")
    lines.append("## 运行统计")
    lines.append("")
    lines.append(f"- 检索数量：{report.total_found}")
    lines.append(f"- 筛选成功：{report.total_filtered}")
    lines.append(f"- AI成功：{report.total_success}")
    lines.append(f"- 失败数量：{report.total_failed}")
    lines.append(f"- API耗时：{report.api_elapsed_seconds:.2f} 秒")
    lines.append("")
    return "\n".join(lines)


def _item_markdown(index: int, item: ReportItem) -> list[str]:
    paper = item.paper
    summary = item.summary
    doi_or_url = paper.url or (f"https://doi.org/{paper.doi}" if paper.doi else "")
    journal_impact = summary.journal_impact or _journal_metadata_text(paper)
    lines = [
        f"### <a id=\"paper-{index}\"></a>{index}. {paper.title} / {summary.chinese_title}",
        "",
        f"- **标题**：{paper.title} / {summary.chinese_title}",
        f"- **作者**：{summary.authors_text or (', '.join(paper.authors[:8]) if paper.authors else '摘要未说明')}",
        f"- **期刊及影响因子**：{journal_impact}",
        f"- **JIF/影响因子**：{paper.impact_factor or '白名单未提供'}",
        f"- **JCR分区**：{paper.quartile or '白名单未提供'}",
        f"- **CiteScore**：{paper.citescore or '白名单未提供'}",
        f"- **出版社**：{paper.publisher or '白名单未提供'}",
        f"- **卷期/DOI**：{summary.volume_issue_doi or paper.doi or '摘要未说明'}",
        f"- **在线发表时间**：{summary.online_date or format_date_cn(paper.published_date)}",
        f"- **SSCI分类**：{'；'.join(paper.ssci_categories) if paper.ssci_categories else 'SSCI白名单匹配；分类未提供'}",
    ]
    if doi_or_url:
        lines.append(f"- **论文链接**：[{doi_or_url}]({doi_or_url})")
    lines.extend(
        [
            "",
            "#### 研究摘要",
            summary.research_abstract,
            "",
            "#### 研究主要内容",
            summary.main_content,
            "",
            "#### 研究论点",
            summary.argument,
            "",
            "#### 研究问题",
            summary.research_question,
            "",
            "#### 研究方法",
            summary.methods,
            "",
            "#### 主要发现",
            summary.findings,
            "",
            "#### 对教师教育研究的启示",
            summary.implications_teacher_education,
            "",
            "#### APA引用格式",
            summary.apa_citation,
            "",
        ]
    )
    return lines


def _journal_metadata_text(paper: Paper) -> str:
    parts = [paper.journal or "未知期刊"]
    if paper.impact_factor:
        parts.append(f"JIF/影响因子：{paper.impact_factor}")
    if paper.quartile:
        parts.append(f"JCR分区：{paper.quartile}")
    if paper.citescore:
        parts.append(f"CiteScore：{paper.citescore}")
    if paper.publisher:
        parts.append(f"出版社：{paper.publisher}")
    if len(parts) == 1:
        parts.append("影响因子/分区：白名单未提供")
    return "；".join(parts)
